"""Render generated artifacts into downloadable documents, in memory.

Two exporters, both returning raw ``bytes`` so the API layer can stream them back
as file downloads:

  * :func:`build_cover_letter_docx` renders the cover letter as a Word ``.docx``
    (via ``python-docx``).
  * :func:`build_resume_pdf` renders the rewritten resume bullets as a ``.pdf``
    (via ``fpdf2``, which is pure-Python and needs no native libraries on any OS).
"""

from io import BytesIO
from typing import Any, Dict, Optional

from docx import Document
from docx.shared import Pt
from fpdf import FPDF
from fpdf.enums import XPos, YPos

# fpdf2's built-in (core) fonts are Latin-1 only, so map the common "smart"
# Unicode characters down to safe ASCII before drawing any PDF text.
_PDF_UNICODE_FIXES = {
    "‘": "'", "’": "'", "“": '"', "”": '"',
    "–": "-", "—": "-", "•": "-", "…": "...",
    " ": " ", "­": "",
}


def _pdf_safe(text: Optional[str]) -> str:
    """Make ``text`` safe to draw with fpdf2's Latin-1 core fonts.

    Folds the common smart punctuation to ASCII and replaces anything else that
    isn't Latin-1, so drawing it never raises.
    """
    if not text:
        return ""
    for unicode_char, ascii_char in _PDF_UNICODE_FIXES.items():
        text = text.replace(unicode_char, ascii_char)
    return text.encode("latin-1", "replace").decode("latin-1")


def build_cover_letter_docx(letter_text: str, *, job_title: str, company: str) -> bytes:
    """Render the cover letter as a Word ``.docx`` document, returned as bytes.

    ``letter_text`` is the letter body as plain / lightly-marked text: blank lines
    separate paragraphs and single newlines become soft line breaks. ``job_title``
    and ``company`` go into a small header line.
    """
    document = Document()

    # Make the default body text a comfortable letter size.
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    header = document.add_paragraph()
    header_run = header.add_run(f"{job_title} — {company}")
    header_run.bold = True

    # Blank lines separate paragraphs; within a paragraph, newlines are soft breaks.
    for block in (letter_text or "").split("\n\n"):
        block = block.strip("\n")
        if not block.strip():
            continue
        paragraph = document.add_paragraph()
        lines = block.split("\n")
        for index, line in enumerate(lines):
            paragraph.add_run(line)
            if index < len(lines) - 1:
                paragraph.add_run().add_break()

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_resume_pdf(resume_rewrite: Optional[Dict[str, Any]], *, job_title: str, company: str) -> bytes:
    """Render the rewritten resume bullets as a ``.pdf``, returned as bytes.

    Bullets are grouped under their section headings, and only the rewritten text
    is exported (this is the polished resume the candidate would actually submit).
    ``job_title`` becomes the title line and ``company`` the subtitle.
    """
    bullets = (resume_rewrite or {}).get("bullets") or []

    pdf = FPDF(format="A4", unit="mm")
    pdf.set_auto_page_break(auto=True, margin=16)
    pdf.set_margins(left=18, top=16, right=18)
    pdf.add_page()

    # Always drop to the next line at the LEFT margin after each block. fpdf2's
    # default leaves the cursor at the right margin, which would zero out the next
    # cell's width.
    def line(height: int, text: str) -> None:
        pdf.multi_cell(0, height, _pdf_safe(text), new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    pdf.set_font("Helvetica", "B", 17)
    line(9, job_title)
    pdf.set_font("Helvetica", "", 12)
    pdf.set_text_color(90, 83, 70)
    line(7, company)
    pdf.set_text_color(20, 20, 20)
    pdf.ln(3)

    current_section = None
    for bullet in bullets:
        section = (bullet.get("section") or "").strip()
        if section and section != current_section:
            current_section = section
            pdf.ln(2)
            pdf.set_font("Helvetica", "B", 12)
            line(7, section)
        pdf.set_font("Helvetica", "", 11)
        line(6, "-  " + (bullet.get("rewritten_bullet") or ""))
        pdf.ln(1)

    if not bullets:
        pdf.set_font("Helvetica", "I", 11)
        line(6, "No resume rewrite was generated.")

    # fpdf2 returns a bytearray; normalize to immutable bytes for the response.
    return bytes(pdf.output())
